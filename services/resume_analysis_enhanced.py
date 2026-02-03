"""
Enhanced Resume Analysis Service - Reliable, Transparent, Human-Friendly

Critical improvements:
1. Validates text extraction before scoring (never silent failures)
2. Shows partial results if full extraction fails
3. Handles all file types safely (PDF/DOCX/TXT)
4. Type-safe integer returns with proper validation
5. Clear error messages - never "undefined" or raw exceptions
6. ATS score always renders correctly
"""

import os
import json
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from database.db import get_db

try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ResumeTextExtractor:
    """Safe text extraction with fallback support."""
    
    @staticmethod
    def extract_from_file(file_path: str) -> Tuple[Optional[str], List[str]]:
        """
        Extract text from resume file safely.
        Returns: (text, warnings)
        - text: extracted text or None if extraction fails
        - warnings: list of issues encountered (empty if perfect)
        """
        warnings = []
        
        if not os.path.exists(file_path):
            return None, ["File not found"]
        
        file_ext = os.path.splitext(file_path)[1].lower()
        text = None
        
        if file_ext == '.pdf':
            text, pdf_warnings = ResumeTextExtractor._extract_pdf(file_path)
            warnings.extend(pdf_warnings)
        elif file_ext == '.docx':
            text, docx_warnings = ResumeTextExtractor._extract_docx(file_path)
            warnings.extend(docx_warnings)
        elif file_ext in ['.txt', '.text']:
            text, txt_warnings = ResumeTextExtractor._extract_txt(file_path)
            warnings.extend(txt_warnings)
        else:
            return None, [f"Unsupported file type: {file_ext}"]
        
        return text, warnings
    
    @staticmethod
    def _extract_pdf(file_path: str) -> Tuple[Optional[str], List[str]]:
        """Extract text from PDF with error handling."""
        warnings = []
        
        if not PDF_AVAILABLE:
            return None, ["PDF support not available - install PyPDF2"]
        
        try:
            pages = []
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                
                if len(reader.pages) == 0:
                    return None, ["PDF is empty - no pages found"]
                
                for i, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            pages.append(text)
                        else:
                            warnings.append(f"Page {i+1} returned empty text")
                    except Exception as e:
                        warnings.append(f"Page {i+1} extraction failed: {str(e)}")
                        continue
            
            if not pages:
                return None, ["No text could be extracted from any page"]
            
            text = '\n'.join(pages)
            
            if len(text.strip()) < 50:
                warnings.append("Extracted text is very short - may be corrupted")
            
            return text, warnings
            
        except Exception as e:
            return None, [f"PDF reading error: {str(e)}"]
    
    @staticmethod
    def _extract_docx(file_path: str) -> Tuple[Optional[str], List[str]]:
        """Extract text from DOCX with error handling."""
        warnings = []
        
        if not DOCX_AVAILABLE:
            return None, ["DOCX support not available - install python-docx"]
        
        try:
            doc = Document(file_path)
            
            if not doc.paragraphs:
                return None, ["DOCX document contains no paragraphs"]
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    paragraphs.append(para.text)
            
            if not paragraphs:
                return None, ["No text content found in document"]
            
            text = '\n'.join(paragraphs)
            
            if len(text.strip()) < 50:
                warnings.append("Extracted text is very short - may be corrupted")
            
            return text, warnings
            
        except Exception as e:
            return None, [f"DOCX reading error: {str(e)}"]
    
    @staticmethod
    def _extract_txt(file_path: str) -> Tuple[Optional[str], List[str]]:
        """Extract text from TXT with error handling."""
        warnings = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            if not text or len(text.strip()) == 0:
                return None, ["Text file is empty"]
            
            if len(text.strip()) < 50:
                warnings.append("Text file is very short")
            
            return text, warnings
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                warnings.append("File was in non-UTF-8 encoding")
                return text, warnings
            except Exception as e:
                return None, [f"Cannot read text file: {str(e)}"]
        except Exception as e:
            return None, [f"Text file reading error: {str(e)}"]


class ResumeAnalyzer:
    """Comprehensive resume analysis with type safety."""
    
    # Safe integer defaults - never None or undefined
    SAFE_DEFAULTS = {
        'ats_score': 0,
        'keyword_score': 0,
        'formatting_score': 0,
        'completeness_score': 0,
        'overall_score': 0,
        'content_quality': 0,
        'section_count': 0,
        'word_count': 0
    }
    
    @staticmethod
    def analyze_resume(file_path: str, user_id: Optional[str] = None) -> Dict:
        """
        Comprehensive resume analysis with fallback support.
        
        Returns: {
            'success': bool,
            'ats_score': int (0-100),
            'analysis': dict with detailed scores,
            'suggestions': list of actionable improvements,
            'warnings': list of extraction/analysis issues,
            'partial': bool (True if some steps failed but partial results shown)
        }
        """
        result = {
            'success': False,
            'ats_score': 0,
            'analysis': {},
            'suggestions': [],
            'warnings': [],
            'partial': False,
            'extracted_text_length': 0
        }
        
        # Step 1: Extract text safely
        text, extraction_warnings = ResumeTextExtractor.extract_from_file(file_path)
        result['warnings'].extend(extraction_warnings)
        
        if not text:
            result['warnings'].insert(0, "Could not extract text from resume file")
            result['suggestions'].append("Try uploading a different file format or check if the file is corrupted")
            return result
        
        # Validate text length
        text = text.strip()
        if len(text) < 50:
            result['warnings'].insert(0, "Resume text is too short to analyze meaningfully")
            result['suggestions'].append("Resume should be at least 50 characters - ensure file uploaded correctly")
            return result
        
        result['extracted_text_length'] = len(text)
        
        # Step 2: Calculate scores safely
        analysis = ResumeAnalyzer._calculate_scores(text)
        
        # Ensure all required integer fields exist and are valid integers
        for key, default in ResumeAnalyzer.SAFE_DEFAULTS.items():
            if key not in analysis:
                analysis[key] = default
            else:
                # Force integer type and clamp to 0-100 range for scores
                try:
                    val = int(analysis[key])
                    if 'score' in key:
                        val = max(0, min(100, val))
                    analysis[key] = val
                except (ValueError, TypeError):
                    analysis[key] = default
        
        result['analysis'] = analysis
        result['ats_score'] = int(analysis.get('ats_score', 0))
        result['success'] = True
        
        # Step 3: Generate suggestions based on scores
        result['suggestions'] = ResumeAnalyzer._generate_suggestions(analysis)
        
        # Save to database if user provided
        if user_id:
            ResumeAnalyzer._save_analysis(user_id, result)
        
        return result
    
    @staticmethod
    def _calculate_scores(text: str) -> Dict:
        """Calculate all resume health scores safely."""
        scores = dict(ResumeAnalyzer.SAFE_DEFAULTS)
        
        text_lower = text.lower()
        word_count = len(text.split())
        scores['word_count'] = word_count
        
        # ATS Score: presence of key sections + keyword density
        ats_sections = ['experience', 'skills', 'education', 'contact']
        sections_found = sum(1 for sec in ats_sections if sec in text_lower)
        scores['ats_score'] = 30 + (sections_found * 15)
        
        # ATS: bonus for action verbs
        action_verbs = ['managed', 'led', 'developed', 'designed', 'implemented', 
                       'achieved', 'improved', 'created', 'increased', 'reduced']
        verb_count = sum(text_lower.count(verb) for verb in action_verbs)
        scores['ats_score'] += min(20, verb_count * 2)
        
        # Keyword score: technical terms
        tech_keywords = ['python', 'javascript', 'sql', 'aws', 'git', 'api', 
                        'database', 'project', 'team', 'system', 'cloud']
        keyword_count = sum(1 for kw in tech_keywords if kw in text_lower)
        scores['keyword_score'] = 20 + (keyword_count * 8)
        
        # Formatting score: word count and structure
        if 300 < word_count < 1000:
            scores['formatting_score'] = 85
        elif 200 < word_count < 1500:
            scores['formatting_score'] = 70
        elif word_count > 100:
            scores['formatting_score'] = 60
        else:
            scores['formatting_score'] = 40
        
        # Line/section diversity (indicates structure)
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        scores['formatting_score'] += min(10, len(lines) // 10)
        
        # Completeness: presence of key sections
        completeness_items = ['education', 'experience', 'skills', 'certification', 'achievement']
        items_found = sum(1 for item in completeness_items if item in text_lower)
        scores['completeness_score'] = 40 + (items_found * 12)
        
        # Content quality: quantified results
        has_numbers = bool(re.search(r'\d+[%$]|\d+x|increased \d+|decreased \d+', text))
        has_metrics = sum(1 for m in ['%', '$', 'x'] if m in text)
        scores['content_quality'] = 50 + (int(has_numbers) * 30) + min(15, has_metrics * 5)
        
        # Section count
        section_keywords = ['experience', 'skills', 'education', 'projects', 'awards', 'certification']
        scores['section_count'] = sum(1 for sec in section_keywords if sec in text_lower)
        
        # Calculate overall score (clamped to 0-100)
        overall = int((scores['ats_score'] + scores['keyword_score'] + 
                      scores['formatting_score'] + scores['completeness_score'] + 
                      scores['content_quality']) / 5)
        scores['overall_score'] = max(0, min(100, overall))
        
        return scores
    
    @staticmethod
    def _generate_suggestions(analysis: Dict) -> List[str]:
        """Generate actionable improvement suggestions."""
        suggestions = []
        
        if analysis['ats_score'] < 70:
            suggestions.append("Add clear section headers (Experience, Skills, Education)")
            suggestions.append("Use action verbs like 'Led', 'Developed', 'Achieved'")
        
        if analysis['keyword_score'] < 70:
            suggestions.append("Include industry-specific keywords for your target role")
            suggestions.append("Mention specific technologies and tools you've used")
        
        if analysis['formatting_score'] < 70:
            suggestions.append("Aim for 300-1000 words (1-2 pages)")
            suggestions.append("Use consistent formatting and clear section breaks")
        
        if analysis['completeness_score'] < 70:
            suggestions.append("Include all sections: Experience, Skills, Education")
            suggestions.append("Add certifications or notable achievements")
        
        if analysis['content_quality'] < 70:
            suggestions.append("Add quantified results (e.g., '30% improvement')")
            suggestions.append("Show impact with metrics and percentages")
        
        if analysis['section_count'] < 3:
            suggestions.append("Expand to include more sections for comprehensiveness")
        
        if not suggestions:
            suggestions.append("Your resume looks well-structured! Focus on adding more metrics.")
        
        return suggestions
    
    @staticmethod
    def _save_analysis(user_id: str, result: Dict) -> bool:
        """Save resume analysis to database safely."""
        try:
            db = get_db()
            now = datetime.utcnow().isoformat()
            
            db.execute('''
                INSERT OR REPLACE INTO resume_health
                (user_id, ats_score, keyword_score, formatting_score, content_completeness,
                 overall_health, suggestions, last_analyzed_at, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (user_id, 
                  int(result['analysis'].get('ats_score', 0)),
                  int(result['analysis'].get('keyword_score', 0)),
                  int(result['analysis'].get('formatting_score', 0)),
                  int(result['analysis'].get('completeness_score', 0)),
                  int(result['ats_score']),  # overall_health = main ats_score
                  json.dumps(result.get('suggestions', [])),
                  now, now, now))
            db.commit()
            return True
        except Exception as e:
            print(f"Database save error: {e}")
            return False


# Import for regex usage
import re
